from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Manual_Completo_Sonnor.docx"


SKIP_DIRS = {
    ".git",
    "node_modules",
    ".expo",
    ".npm-cache",
    ".pap_render",
    "tmp",
    "__pycache__",
}


def rel(path: Path) -> str:
    return str(path.relative_to(ROOT)).replace("\\", "/")


def list_tree(path: Path, prefix: str = "", depth: int = 0, max_depth: int = 6) -> list[str]:
    if depth > max_depth:
        return []
    entries = sorted(
        [p for p in path.iterdir() if p.name not in SKIP_DIRS],
        key=lambda p: (not p.is_dir(), p.name.lower()),
    )
    lines: list[str] = []
    for index, entry in enumerate(entries):
        connector = "`-- " if index == len(entries) - 1 else "|-- "
        lines.append(f"{prefix}{connector}{entry.name}{'/' if entry.is_dir() else ''}")
        if entry.is_dir():
            extension = "    " if index == len(entries) - 1 else "|   "
            lines.extend(list_tree(entry, prefix + extension, depth + 1, max_depth))
    return lines


def read_text(path: str) -> str:
    return (ROOT / path).read_text(encoding="utf-8", errors="ignore")


def extract_exports(path: str) -> list[str]:
    text = read_text(path)
    exports = []
    for line in text.splitlines():
        match = re.match(r"\s*export\s+(?:async\s+)?(?:function|const|default function)\s+([A-Za-z0-9_]+)", line)
        if match:
            exports.append(match.group(1))
    return exports


def line_count(path: Path) -> int:
    return len(path.read_text(encoding="utf-8", errors="ignore").splitlines())


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "F2F4F7")
        set_cell_text(cell, header, bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            cell.width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()
    return table


def add_code_block(doc: Document, text: str, max_chars: int | None = None):
    if max_chars and len(text) > max_chars:
        text = text[:max_chars].rstrip() + "\n... (continua no projeto)"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(7.5)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def package_rows() -> list[list[str]]:
    package = json.loads(read_text("package.json"))
    rows = []
    explanations = {
        "@expo/vector-icons": "Biblioteca de ícones pronta para Expo. Serve para usar ícones visuais em botões, menus e ações sem desenhar tudo manualmente.",
        "@react-native-async-storage/async-storage": "Armazenamento local simples no telemóvel. Na Sonnor é usado principalmente para o Firebase Auth guardar a sessão do utilizador entre aberturas da app.",
        "@react-navigation/bottom-tabs": "Navegação por abas inferiores. Mesmo usando Expo Router, esta biblioteca dá suporte à navegação tipo tab bar em React Native.",
        "@react-navigation/elements": "Componentes internos de navegação, headers e elementos usados pelo ecossistema React Navigation/Expo Router.",
        "@react-navigation/native": "Base de navegação React Navigation. O Expo Router usa esta fundação para controlar rotas, transições e histórico.",
        "expo": "Base da app mobile; permite correr em Android, iOS e Web.",
        "expo-av": "Áudio principal do player. É o pacote que permite carregar MP3, tocar, pausar, avançar, voltar, procurar posição e receber progresso da música.",
        "expo-blur": "Cria fundos desfocados. Serve para painéis visuais, menus e elementos sobrepostos com aparência moderna.",
        "expo-constants": "Dá acesso a constantes do ambiente Expo, como dados da app, plataforma e configuração em runtime.",
        "expo-document-picker": "Permite escolher documentos/ficheiros do dispositivo. Na Sonnor é importante para escolher ficheiros MP3 no envio de músicas.",
        "expo-font": "Carrega fontes personalizadas. Na Sonnor carrega a fonte Bristol usada na abertura e identidade visual.",
        "expo-haptics": "Permite vibrações/feedback tátil em toques. Serve para dar sensação física em ações importantes, botões e interações.",
        "expo-image": "Componente otimizado de imagem. Serve para mostrar capas, avatars, banners, posts e imagens com melhor cache/performance que Image simples.",
        "expo-image-picker": "Abre galeria/câmara para escolher imagens e vídeos. Usado para avatar, banner, posts, capas e eventos.",
        "expo-linking": "Trabalha com links e deep links. Ajuda a app a abrir rotas internas ou URLs externas quando necessário.",
        "expo-media-library": "Acesso à biblioteca de media do dispositivo. Ajuda em fluxos ligados a imagens/vídeos locais.",
        "expo-router": "Sistema de navegação por ficheiros dentro da pasta app. Cada ficheiro em app/ vira uma rota/tela.",
        "expo-splash-screen": "Controla o ecrã inicial nativo. Na Sonnor evita esconder o splash antes da fonte, Firebase e animação de abertura estarem prontos.",
        "expo-status-bar": "Controla a barra de estado do telemóvel. Na Sonnor usa estilo claro sobre fundo escuro.",
        "expo-symbols": "Fornece símbolos nativos/ícones em ambientes compatíveis, útil para UI com aspeto mais nativo.",
        "expo-system-ui": "Controla cores e comportamento da UI do sistema. Ajuda a manter a experiência escura/edge-to-edge coerente.",
        "expo-video": "Reprodução de vídeo moderna do Expo. Usado em posts de vídeo, vídeos curtos de músicas e prévias fullscreen.",
        "expo-web-browser": "Abre páginas externas no browser seguro do sistema. Útil para links externos como Spotify, lojas, eventos ou páginas web.",
        "firebase": "SDK do Firebase no cliente. Dá Auth, Firestore, Storage e chamadas a Cloud Functions usadas por quase toda a app.",
        "react": "Biblioteca base para construir componentes. A Sonnor inteira é feita com componentes React.",
        "react-dom": "Renderização React para Web. Necessário quando a app Expo também corre no navegador.",
        "react-native": "Framework mobile que transforma componentes React em interface nativa Android/iOS.",
        "react-native-gesture-handler": "Gestos avançados: arrastar, tocar, swipes e interações complexas. Essencial para media fullscreen, sliders e menus.",
        "react-native-reanimated": "Animações fluidas no lado nativo. Ajuda em transições, microinterações e movimentos sem travar a interface.",
        "react-native-safe-area-context": "Respeita áreas seguras do telemóvel, como notch, barra inferior e cantos. Evita conteúdo tapado.",
        "react-native-screens": "Melhora performance das telas nativas usadas pela navegação, mantendo rotas mais eficientes.",
        "react-native-svg": "Permite SVGs e ícones vetoriais. Usado para ícones próprios e elementos gráficos escaláveis.",
        "react-native-web": "Permite que componentes React Native rodem no navegador quando se usa Expo Web.",
        "react-native-worklets": "Suporte de worklets usado por animações/gestos modernos, especialmente no ecossistema Reanimated.",
        "@types/node": "Tipos TypeScript para APIs de Node. Ajuda scripts e tooling a reconhecerem objetos/funções de Node.",
        "@types/react": "Tipos TypeScript do React. Dá autocomplete e validação de props, hooks e componentes.",
        "eslint": "Ferramenta de lint. Procura problemas de código, padrões perigosos e inconsistências.",
        "eslint-config-expo": "Configuração ESLint recomendada para projetos Expo, alinhada com React Native e Expo Router.",
        "typescript": "Linguagem usada no projeto para adicionar tipos ao JavaScript. Ajuda a evitar erros em telas, Firebase e dados.",
    }
    for name, version in package["dependencies"].items():
        rows.append([name, version, explanations[name]])
    for name, version in package["devDependencies"].items():
        rows.append([name, version, explanations[name]])
    return rows


ROUTE_EXPLANATIONS = [
    ["app/_layout.tsx", "Entrada global", "Carrega fonte Bristol, mostra splash preto com S, prepara Firebase e envolve a app com PlayerProvider e SuccessFeedbackProvider."],
    ["app/index.tsx", "Decisão inicial", "Verifica sessão e envia o utilizador para login, onboarding ou área principal."],
    ["app/auth/login.tsx", "Login", "Formulário de email e senha, validações, alternância de ver senha e entrada no Firebase Auth."],
    ["app/auth/new-account.tsx", "Criação de conta", "Fluxo de registo com código por email antes de criar a conta."],
    ["app/auth/verify-mail.tsx", "Verificação", "Confirma o código enviado por email durante o registo."],
    ["app/auth/forgot-password.tsx", "Recuperação", "Pede código OTP para redefinir password."],
    ["app/auth/new-password.tsx", "Nova password", "Confirma token/código e chama a Cloud Function para trocar a senha."],
    ["app/onboarding/create-profile.tsx", "Onboarding", "Primeira configuração do perfil: nome, username, dados visuais e continuação para a app."],
    ["app/main/_layout.tsx", "Layout principal", "Controla quando aparece o menu inferior dinâmico."],
    ["app/main/home.tsx", "Home", "Feed principal: músicas, posts, eventos, pré-lançamentos, recomendações e player integrado."],
    ["app/main/search.tsx", "Pesquisa/Explorar", "Pesquisa músicas, perfis, albums e posts; inclui grelha visual e visualização fullscreen."],
    ["app/main/library.tsx", "Biblioteca", "Mostra músicas gostadas, albums guardados, playlists e histórico recente."],
    ["app/main/profile.tsx", "Perfil", "Perfil público/privado com posts, músicas, albums, merchandising, follows e ações de dono."],
    ["app/main/profile/edit-profile.tsx", "Editar perfil", "Atualiza dados pessoais, avatar, banner, links e interesses."],
    ["app/main/profile/organize.tsx", "Organização", "Permite organizar conteúdo do perfil."],
    ["app/main/create/track.tsx", "Enviar música", "Escolhe MP3, declara titularidade e cria submissões de música para revisão."],
    ["app/main/create/finish-track/[submissionId].tsx", "Finalizar música", "Depois da aprovação, adiciona capa, dados, letra, género, vídeo curto e cria faixa/album."],
    ["app/main/events/request.tsx", "Pedido de evento", "Utilizador envia banner de evento para aprovação."],
    ["app/main/release/[slug].tsx", "Página de lançamento", "Mostra album/single/EP, capa, fundo, lista de faixas, contador de pré-lançamento e play queue."],
    ["app/main/track/[id].tsx", "Detalhe da faixa", "Redireciona para o lançamento ou contexto da música."],
    ["app/main/post/[id].tsx", "Detalhe de post", "Abre conteúdo específico de post."],
    ["app/main/settings/account.tsx", "Conta", "Pequena tela de configurações de conta."],
    ["app/admin/index.tsx", "Admin geral", "Painel administrativo com atalhos, métricas e estado da plataforma."],
    ["app/admin/music-reviews.tsx", "Revisão musical", "Admin aprova/rejeita submissões, pede provas e controla riscos de direitos."],
    ["app/admin/manage-users.tsx", "Utilizadores", "Admin vê perfis, verificação e ações sensíveis como eliminação."],
    ["app/admin/user-requests.tsx", "Pedidos", "Admin gere pedidos de perfil/conteúdo."],
    ["app/admin/event-banners.tsx", "Eventos", "Admin aprova pedidos de evento e publica banners."],
]


DATA_MODEL = [
    ["users", "Perfil do utilizador", "username, displayName, bio, avatar/banner/background, merchandising, cidade/país, verificação, contadores."],
    ["albums", "Lançamentos/pastas", "title, slug, type album/single/ep, capa, fundo, releaseDate, pré-lançamento, géneros, trackIds e contadores."],
    ["tracks", "Músicas", "audioUrl, previewUrl, shortVideoUrl, letras, featUserIds, género, status, albumId, likes/plays/comments."],
    ["musicSubmissions", "Submissões para revisão", "MP3 original, declaração de direitos, estado de revisão, fingerprint e decisão da equipa."],
    ["posts", "Publicações sociais", "caption, mediaUrl, image/video, overlays, música ligada, clip start/end, categoria e contadores."],
    ["comments", "Comentários", "targetType, targetId, userId, text, likesCount."],
    ["likes", "Gostos", "Documento por user + target, criado por Cloud Function para manter contadores certos."],
    ["follows", "Seguidores", "Relação followerId/followingId e subcoleções users/{id}/followers/following."],
    ["playlists", "Playlists", "Nome, capa, público/privado, trackIds."],
    ["collections", "Coleções", "Guarda tracks, posts e albums juntos."],
    ["recentPlays", "Histórico", "trackId, albumId, listenedMs, completed e origem."],
    ["notifications", "Notificações", "Tipo, título, corpo, target e estado read."],
    ["releaseReminders", "Lembretes", "Utilizador pede aviso quando pré-lançamento fica disponível."],
    ["messageThreads/messages", "Mensagens", "Estrutura preparada para conversas, mesmo que muitas telas tenham sido removidas."],
    ["reports", "Reports", "Denúncias de user/post/track/album/comment com resposta admin."],
    ["eventRequests/eventBanners", "Eventos", "Pedido do utilizador e banner aprovado publicado por admin."],
    ["profileRequests", "Pedidos de perfil/conteúdo", "Mudança de nome, pedidos de remoção, reports e decisões admin."],
    ["verificationRequests", "Verificação", "Provas enviadas e estado pending/approved/rejected."],
    ["appConfig", "Configuração pública/admin", "Destaques, categorias e lista de admins."],
]


CLOUD_FUNCTIONS = [
    ["sendOtpEmail / verifyOtp / resetPasswordWithOtp", "Recuperação de password por código de 4 dígitos, com hash, expiração e limite de tentativas."],
    ["sendSignupOtpEmail / verifySignupOtp", "Código de criação de conta antes do registo."],
    ["sendSignupOtpHttp / verifySignupOtpHttp", "Fallback HTTP quando callable não responde em certos ambientes."],
    ["sendDeleteAccountOtpEmail / deleteAccountWithOtp", "Confirma eliminação de conta por código e depois apaga dados ligados."],
    ["deleteAccountWithOtpHttp", "Fallback HTTP para eliminação de conta."],
    ["getCurrentAdminStatus", "Verifica se o utilizador é admin por custom claim ou email em appConfig/admins."],
    ["sendMusicOwnershipContactEmail", "Email bonito para pedir prova de titularidade de músicas."],
    ["adminDeleteUserAccount", "Admin elimina uma conta sem precisar do código do próprio utilizador."],
    ["setUserVerifiedOverride", "Admin força ou remove verificação manual."],
    ["toggleTrackLikeV2 / togglePostLikeV2 / toggleContentLike", "Gosto/desgosto transacional com atualização de contadores."],
    ["setUserFollow", "Follow/unfollow transacional, cria relação e atualiza contadores."],
    ["deleteOwnedAlbum", "Dono apaga album e dados ligados com limpeza."],
    ["publishScheduledReleases", "Agenda a cada 1 minuto para publicar lançamentos quando a data chega."],
    ["expireEventBanners", "Agenda diária que marca banners expirados."],
    ["publishDueReleasesNow", "Força publicação de lançamento que já chegou à data."],
    ["notifyMusicSubmissionDecision", "Cria notificação quando submissão musical muda de decisão."],
    ["notifyEventRequestDecision", "Cria notificação quando evento é aprovado/rejeitado."],
    ["notifyReportDecision", "Cria notificação quando admin responde report."],
    ["notifyVerificationDecision", "Cria notificação quando verificação muda de estado."],
    ["notifyProfileRequestDecision", "Cria notificação para pedidos de perfil/conteúdo."],
    ["recalculateVerificationOnUserWrite/TrackWrite/ReportWrite", "Recalcula verificação automática por seguidores, hits e reports."],
    ["processAdminDeletionRequest", "Processa pedidos administrativos de eliminação criados no Firestore."],
]


CLIENT_EXPLANATIONS = [
    ["firebase/config.ts", "Inicializa Firebase, Auth com AsyncStorage, Functions e exporta app/auth/functions."],
    ["firebase/dataClient.ts", "Exporta Firestore e Storage a partir da app inicializada."],
    ["firebase/auth.ts", "Login, registo, logout, recuperação de password, OTP de signup e eliminação de conta."],
    ["firebase/userProfile.ts", "Cria perfil padrão se não existir e atualiza dados do perfil."],
    ["firebase/contentClient.ts", "Leitura/caching de home, search, library, profile, album, track e recomendações."],
    ["firebase/contentMutations.ts", "Criação/atualização de tracks, albums, posts, collections e recent plays."],
    ["firebase/socialClient.ts", "Likes, follows, comentários, reports, verification requests e notificações."],
    ["firebase/storageClient.ts", "Upload resumível para Storage, tipos de ficheiro e cache busting."],
    ["firebase/musicReviewClient.ts", "Admin lê/aprova/rejeita submissões musicais e contacta artistas."],
    ["firebase/eventClient.ts", "Criação de pedidos de evento e gestão/admin de banners."],
    ["firebase/profileRequests.ts", "Pedidos de mudança/remoção/report e decisão admin."],
    ["firebase/adminClient.ts", "Operações administrativas: overview, admins, eliminação, verificação."],
    ["firebase/settingsClient.ts", "Settings do utilizador, histórico de pesquisa, bloquear utilizadores e biblioteca."],
    ["firebase/paths.ts", "Todos os nomes de coleções e caminhos de Firestore/Storage centralizados."],
    ["firebase/schema.ts", "Tipos TypeScript que documentam a base de dados."],
]


GIT_COMMANDS = [
    ["git --version", "Confirma se o Git está instalado no computador."],
    ["git init", "Cria um repositório Git numa pasta nova. Só se usa uma vez quando o projeto ainda não tem Git."],
    ["git status", "Mostra ficheiros alterados, novos ou apagados. É o comando mais usado antes de guardar alterações."],
    ["git add nome-do-ficheiro", "Prepara um ficheiro específico para entrar no próximo commit."],
    ["git add .", "Prepara todas as alterações da pasta atual. Usa com cuidado para não incluir ficheiros gerados sem querer."],
    ["git commit -m \"mensagem\"", "Guarda uma versão do projeto no histórico local com uma mensagem explicativa."],
    ["git log --oneline", "Mostra o histórico de commits de forma curta."],
    ["git branch", "Mostra as branches existentes e indica em qual estás."],
    ["git branch nome", "Cria uma branch nova para trabalhar sem mexer diretamente na main."],
    ["git switch nome", "Muda para outra branch."],
    ["git switch -c nome", "Cria e muda para uma branch nova no mesmo comando."],
    ["git remote -v", "Mostra para que GitHub o projeto está conectado."],
    ["git remote add origin URL", "Liga um projeto local a um repositório GitHub quando ainda não existe remote."],
    ["git push -u origin main", "Envia a branch main para o GitHub pela primeira vez e grava o tracking."],
    ["git push", "Envia commits locais para o GitHub depois da primeira configuração."],
    ["git pull", "Baixa alterações do GitHub e junta com a tua versão local."],
    ["git clone URL", "Copia um projeto do GitHub para o computador."],
    ["git diff", "Mostra linha por linha o que mudou antes do commit."],
    ["git restore ficheiro", "Descarta alterações locais de um ficheiro. Usar só quando tens certeza."],
]


CODE_BASICS = [
    ["const", "Cria uma constante: um nome que não pode receber outro valor depois. Exemplo: const email = 'a@b.com'. O conteúdo de objetos/arrays ainda pode mudar, mas a variável não aponta para outro objeto."],
    ["let", "Cria uma variável que pode mudar. Exemplo: let loading = true; depois loading = false."],
    ["function", "Define um bloco de código reutilizável. Exemplo: function normalizeEmail(email) { return email.trim().toLowerCase(); }."],
    ["async function", "Função que faz operações assíncronas, como ler Firebase, enviar email ou fazer upload. Normalmente usa await dentro."],
    ["await", "Espera uma promessa terminar antes de continuar. Na Sonnor aparece em chamadas ao Firebase, uploads e Cloud Functions."],
    ["import", "Traz código de outro ficheiro ou biblioteca. Exemplo: import { auth } from '../firebase/config'."],
    ["export", "Permite que uma função, constante ou tipo seja usado por outros ficheiros."],
    ["export default", "Exporta o item principal do ficheiro. Em telas Expo Router normalmente é o componente da tela."],
    ["type", "Cria um tipo TypeScript. Não existe no app final; serve para o editor/compilador saber o formato dos dados."],
    ["Props", "Dados recebidos por um componente React. Exemplo: um componente pode receber title, uri ou onPress."],
    ["useState", "Hook React para guardar estado da tela. Exemplo: const [loading, setLoading] = useState(false)."],
    ["useEffect", "Hook React para executar código quando a tela abre, quando um valor muda ou quando precisa limpar algo."],
    ["useRef", "Guarda uma referência que sobrevive entre renders sem redesenhar a tela. Usado para players, timers e valores internos."],
    ["useMemo", "Calcula um valor e reaproveita enquanto as dependências não mudam. Ajuda performance."],
    ["useCallback", "Guarda uma função entre renders quando ela depende de certos valores. Ajuda performance e evita recriar funções."],
    ["return (...)", "Em componentes React, devolve o visual que aparece no ecrã."],
    ["StyleSheet.create", "Cria estilos React Native parecidos com CSS, mas em objeto JavaScript."],
    ["onPress", "Função executada quando o utilizador toca num botão/elemento."],
    ["try/catch", "Tenta executar código e captura erro se falhar. Usado em login, uploads, Firebase e Functions."],
    ["Promise", "Representa uma operação que ainda não terminou, como download, upload ou leitura do Firestore."],
    ["array .map", "Transforma uma lista noutra lista. Exemplo: tracks.map(track => mostrar uma linha por música)."],
    ["array .filter", "Filtra uma lista. Exemplo: posts.filter(post => post.status === 'published')."],
    ["object", "Estrutura com chaves e valores. Exemplo: { title: 'Música', likesCount: 0 }."],
    ["template string", "Texto com variáveis usando crases. Exemplo: `users/${userId}/avatar.jpg`."],
    ["?.", "Optional chaining. Acede a algo só se existir. Exemplo: user?.uid evita erro quando user é null."],
    ["??", "Valor padrão quando algo é null ou undefined. Exemplo: input.status ?? 'draft'."],
]


SCREEN_ROUTING_LESSONS = [
    ["app/index.tsx", "É a primeira tela visível depois do layout global. Se já houver user logado, usa router.replace('/main/home'). Se não houver, mostra botões Entrar e Criar Conta."],
    ["app/auth/login.tsx", "É chamada quando o utilizador toca em Entrar. A rota vem do caminho do ficheiro: /auth/login."],
    ["app/auth/new-account.tsx", "É chamada quando o utilizador toca em Criar Conta. A rota é /auth/new-account."],
    ["app/main/home.tsx", "É a home principal. A rota é /main/home. Mostra feed, músicas, posts, eventos e usa dados do Firebase."],
    ["app/main/_layout.tsx", "É o layout das telas main. Ele renderiza a Stack e decide se o Dynamicmenu aparece."],
    ["app/main/release/[slug].tsx", "O [slug] significa rota dinâmica. Exemplo: /main/release/meu-album abre a mesma tela, mas com slug diferente."],
    ["app/main/track/[id].tsx", "O [id] também é rota dinâmica. Serve para abrir detalhes de uma música específica."],
    ["app/main/create/finish-track/[submissionId].tsx", "Recebe o ID da submissão aprovada pela URL e carrega os dados para finalizar a música."],
]


VISUAL_BUILDING_BLOCKS = [
    ["View", "É como uma div/bloco. Serve para agrupar conteúdo, criar linhas, colunas, fundos e áreas."],
    ["Text", "Mostra texto. Todo texto visível no React Native precisa estar dentro de Text."],
    ["Image", "Mostra imagem local ou remota: avatar, capa, banner, GIF ou post."],
    ["ImageBackground", "Imagem usada como fundo de uma área, comum em capas e banners."],
    ["ScrollView", "Cria uma tela que pode rolar verticalmente/horizontalmente."],
    ["Pressable / TouchableOpacity", "Elementos clicáveis. Usados para botões, cards tocáveis, ações de menu e navegação."],
    ["Modal", "Janela por cima da tela atual. Usada para menus, confirmações, reports e painéis."],
    ["BlurView", "Camada com desfoque. Ajuda a criar visual moderno em menus sobrepostos."],
    ["StyleSheet.create", "Onde ficam os estilos: cor, tamanho, margem, padding, flex, borderRadius, posição, etc."],
    ["Ionicons", "Ícones prontos usados em botões, setas, menu, play/pause, perfil e ações."],
]


MENU_LESSONS = [
    ["Onde está", "O menu principal está em app/main/components/Dynamicmenu.tsx."],
    ["Quem chama o menu", "app/main/_layout.tsx importa Dynamicmenu e mostra o componente no fim da tela quando showDynamic é true."],
    ["Quando aparece", "Aparece em /main/home, /main/search, /main/profile, /main/library e rotas que começam com /main/release/."],
    ["Como sabe a tela atual", "usePathname() lê a rota atual. Depois o código compara o pathname com as rotas permitidas."],
    ["Como abre telas", "Dentro do menu usa useRouter() e chamadas como router.push('/main/profile') ou router.replace('/auth/login')."],
    ["Como abre popups", "Usa useState para controlar booleanos como menuVisible, userMenuVisible, privacyVisible, statsVisible e reportVisible. Quando true, o Modal aparece."],
    ["Como mostra o mini player", "Dynamicmenu usa usePlayer() para receber track, status, togglePlay e seek. Assim o menu consegue mostrar música atual e controlar play/pause/progresso."],
    ["Como carrega dados do perfil", "Usa useCurrentUser() para saber quem está logado e getProfileContent(user.uid) para buscar avatar, username, seguidores e músicas."],
    ["Como sabe se é admin", "Chama isCurrentUserAdmin() e guarda em isAdmin. Se for admin, mostra opções administrativas."],
]


VISUAL_PATTERNS = [
    ["Fundo escuro", "A maioria das telas usa backgroundColor '#000'. Isso dá identidade musical, foco em media e contraste com capas."],
    ["Cards/blocos", "Blocos usam borderRadius, padding, backgroundColor rgba(...) e borderColor suave. Exemplo: AppScreen cria cards reutilizáveis para seções."],
    ["Botões", "Botões usam Pressable/TouchableOpacity, borderRadius alto, alinhamento central e texto em branco/preto conforme o fundo."],
    ["Layout responsivo", "utils/responsive.ts fornece wp, hp e font para adaptar tamanhos conforme largura/altura do dispositivo."],
    ["Safe area visual", "AppScreen e layouts dão paddingTop/paddingBottom para não colar no topo, notch ou menu."],
    ["Media grande", "Home, Search e Profile usam capas, posts e vídeos como foco visual, não apenas texto."],
    ["Overlays", "Posts podem ter overlayMedia: imagens por cima da media base com x, y, scale e dimensões de palco."],
    ["Fullscreen", "FullscreenMedia.tsx e fullmidia.tsx ampliam imagens/vídeos para experiência imersiva."],
    ["Feedback", "pressFeedback e SuccessFeedback tornam toques e ações mais visíveis para o utilizador."],
]


SCREEN_CREATION_STEPS = [
    "Criar um ficheiro dentro de app/. O caminho do ficheiro vira a rota automaticamente.",
    "Exportar uma função default, por exemplo: export default function MinhaTela() { ... }.",
    "Importar componentes do React Native, como View, Text, Pressable e StyleSheet.",
    "Usar useRouter() se a tela precisar navegar para outra rota.",
    "Usar useState para guardar campos, loading, erros, popups e seleções.",
    "Usar useEffect para carregar dados quando a tela abre ou quando o user muda.",
    "Chamar funções da pasta firebase/ para ler ou gravar dados.",
    "Retornar o JSX com o visual da tela.",
    "Criar styles no fim com StyleSheet.create.",
]


CODE_FILE_LESSONS = [
    ["app/_layout.tsx", "É o cérebro inicial visual. Carrega fonte Bristol, controla splash, espera Firebase e coloca providers globais. Sem ele, a app não teria abertura nem contexto global."],
    ["app/main/_layout.tsx", "É o layout da área logada. Mostra a Stack de telas e põe Dynamicmenu por cima quando a rota permite."],
    ["components/AppScreen.tsx", "É um molde reutilizável de tela: recebe title, subtitle, actions e sections. Depois monta header, botões e cards automaticamente."],
    ["app/main/components/Dynamicmenu.tsx", "É o menu dinâmico e mini player. Tem estados para popups, perfil, reports, privacidade, estatísticas, delete account e player."],
    ["app/main/home.tsx", "É uma tela grande porque junta muitos blocos: banners, posts, músicas, pré-lançamentos, clips de áudio, likes, reports e navegação."],
    ["app/main/search.tsx", "Mostra pesquisa e exploração. Usa filtros, grelha de posts, resultados de users/tracks/albums e visualização fullscreen."],
    ["app/main/profile.tsx", "Mostra perfil completo. Lida com avatar, banner, conteúdo, follows, posts, músicas, albums, loja/merch e ações do dono."],
    ["context/PlayerContext.tsx", "Centraliza o player. Qualquer tela chama usePlayer() e consegue tocar músicas sem recriar a lógica."],
    ["firebase/contentClient.ts", "Centraliza leituras e cache. As telas pedem getHomeContent, getProfileContent ou getAlbumContent em vez de repetir queries Firestore."],
    ["firebase/contentMutations.ts", "Centraliza escritas de conteúdo: criar música, album, post, collection e histórico."],
    ["firebase/socialClient.ts", "Centraliza ações sociais: likes, follows, comentários, reports, notificações e verificação."],
    ["functions/src/index.ts", "Backend seguro. Faz OTP, emails, admin, likes/follows transacionais, publicação agendada e notificações automáticas."],
]


DEEP_SCREEN_LESSONS = [
    ["app/index.tsx", "Mostra a primeira experiência: logo Sonnor, GIF motion, slogan e botões Entrar/Criar Conta. Usa useCurrentUser para saber se já existe sessão. Se existir, redireciona automaticamente para /main/home."],
    ["app/auth/login.tsx", "Controla email, password, loading, ver/ocultar password e erros. Quando o utilizador toca em Entrar, chama login() em firebase/auth.ts. Se der certo, a sessão fica guardada pelo Firebase Auth."],
    ["app/auth/new-account.tsx + verify-mail.tsx", "Implementa registo com código. Primeiro pede email, chama sendSignupCode, depois verifySignupCode confirma OTP, e só depois cria a conta com password."],
    ["app/auth/forgot-password.tsx + new-password.tsx", "Fluxo de recuperação: pede código por email, verifica o código, recebe resetToken e troca a password por Cloud Function."],
    ["app/onboarding/create-profile.tsx", "Tela para completar o perfil inicial. Recolhe dados como nome, username, bio e imagem. Depois grava em users/{uid}."],
    ["app/main/home.tsx", "Uma das telas principais. Carrega getHomeContent, monta listas de músicas, albums, posts, eventos, pré-lançamentos e recomendações. Também permite tocar música, abrir release, dar like, reportar e ver posts com clips de áudio."],
    ["app/main/search.tsx", "Tela de pesquisa/exploração. Normaliza texto pesquisado, mistura resultados de músicas, albums, perfis e posts, mostra histórico de pesquisa e abre posts em sequência fullscreen."],
    ["app/main/library.tsx", "Mostra biblioteca pessoal: músicas gostadas, albums guardados, histórico recente e playlists. Dados vêm de getLibraryContent."],
    ["app/main/profile.tsx", "Tela mais completa de perfil. Mostra avatar, banner, bio, músicas, albums, posts, loja/merch, seguidores, seguir/deixar de seguir, edição quando é o próprio perfil e ações de report."],
    ["app/main/profile/edit-profile.tsx", "Edita dados do perfil. Pode atualizar campos e imagens, usando updateUserProfile e uploads quando necessário."],
    ["app/main/create/track.tsx", "Começa publicação musical. Escolhe MP3, valida formato, recolhe declaração de titularidade, define batch/single/EP/album e cria musicSubmissions para revisão."],
    ["app/main/create/finish-track/[submissionId].tsx", "Finaliza música já aprovada. Carrega submission, permite adicionar capa, título, género, letras, feat, vídeo curto, data de lançamento e cria track/album final."],
    ["app/main/components/PopUpCreate/createPost.tsx", "Editor de post. Escolhe imagem/vídeo, escala media, adiciona overlays, associa música, escolhe trecho do áudio e publica o post no Firestore/Storage."],
    ["app/main/release/[slug].tsx", "Página do lançamento. Mostra capa/fundo, lista de faixas, status de pré-lançamento, contador, botão play e lembrete."],
    ["app/admin/music-reviews.tsx", "Admin vê submissões musicais, agrupa por batch, aprova, rejeita, pede prova de titularidade e marca risco/revisão."],
    ["app/admin/event-banners.tsx", "Admin vê pedidos de evento, aprova para gerar banner publicado por 7 dias, rejeita com motivo ou remove banners."],
    ["app/admin/manage-users.tsx", "Admin lista utilizadores, vê verificação, seguidores/músicas, força verificação ou pede eliminação."],
]


FIREBASE_CLIENT_DEEP = [
    ["firebase/config.ts", "Cria o app Firebase com apiKey, projectId e storageBucket. Inicializa Auth com persistência AsyncStorage para a sessão não sumir ao fechar a app. Exporta auth e functions."],
    ["firebase/dataClient.ts", "Ponto simples que exporta db = getFirestore(app) e storage = getStorage(app). Outros ficheiros importam daqui."],
    ["firebase/paths.ts", "Evita escrever nomes de coleções/caminhos espalhados. Exemplo: firestorePaths.user(userId) devolve users/{userId}; storagePaths.trackAudio(trackId) devolve tracks/{trackId}/audio.mp3."],
    ["firebase/schema.ts", "Define os formatos dos documentos: UserDocument, TrackDocument, AlbumDocument, PostDocument, MusicSubmissionDocument, etc. Ajuda a saber que campos existem."],
    ["firebase/auth.ts", "Liga telas de auth às Cloud Functions. Normaliza email, faz login/register/logout, envia OTP, verifica OTP, redefine senha e apaga conta com código."],
    ["firebase/userProfile.ts", "Garante que todo utilizador Auth tem documento em users. Se não existir, cria perfil inicial com username privado, verified false e contadores zero."],
    ["firebase/contentClient.ts", "Camada de leitura. Lê coleções com fallback/cache, monta dados compostos da home, perfil, album, library e search. Evita repetir queries nas telas."],
    ["firebase/contentMutations.ts", "Camada de escrita de conteúdo. Cria tracks, albums, posts, collections e recentPlays; atualiza media/detalhes e apaga posts com comentários/likes ligados."],
    ["firebase/storageClient.ts", "Transforma URI local em Blob, escolhe caminho certo no Storage, faz uploadBytesResumable, calcula progresso e devolve downloadUrl."],
    ["firebase/socialClient.ts", "Ações sociais. Likes e follows vão para Cloud Functions com idToken; comentários/reports/verification requests são gravados no Firestore."],
    ["firebase/musicReviewClient.ts", "Submissão e revisão musical. Cria musicSubmissions, faz upload MP3, lista pendentes, aprova/rejeita, permite batch e chama email de titularidade."],
    ["firebase/eventClient.ts", "Eventos. Cria eventRequests com upload de banner, lista pedidos, aprova gerando eventBanners por 7 dias, rejeita e remove banners."],
    ["firebase/adminClient.ts", "Admin. Lista users/posts/releases/reports, verifica se user é admin, chama setUserVerifiedOverride e cria adminDeletionRequests."],
]


ALL_FIREBASE_FILES = [
    ["adminClient.ts", "Código do admin no app. Lista utilizadores, posts, releases, reports e pedidos; verifica se o user atual é admin; chama Cloud Functions para verificação manual; cria pedido para apagar utilizador."],
    ["albumDeletionClient.ts", "Cliente pequeno para apagar album do próprio utilizador. Pega idToken, chama Cloud Function deleteOwnedAlbum e limpa cache de conteúdo."],
    ["auth.ts", "Tudo de autenticação: register, login, logout, reset password por OTP, OTP de signup, OTP para apagar conta, fallback HTTP e normalização de email."],
    ["config.ts", "Configuração central do Firebase: apiKey, projectId, storageBucket, initializeApp, Auth com AsyncStorage e Functions."],
    ["contentClient.ts", "Leituras principais da app. Busca home, profile, library, album, track, recomendações, users reportáveis e usa cache para evitar repetir pedidos."],
    ["contentMutations.ts", "Escritas principais de conteúdo. Cria/atualiza tracks, albums, posts, collections, recentPlays e apaga posts com likes/comentários ligados."],
    ["dataClient.ts", "Exporta db e storage. É o atalho para Firestore e Firebase Storage usado pelos outros ficheiros."],
    ["defaultContent.ts", "Define dados vazios/padrão para quando ainda não há conteúdo ou quando uma leitura falha: user padrão, listas vazias de tracks, releases, posts e boxes."],
    ["eventClient.ts", "Fluxo de eventos. Utilizador pede banner, upload vai para Storage, admin lista pedidos, aprova criando eventBanner, rejeita ou remove banner."],
    ["musicReviewClient.ts", "Fluxo de revisão musical. Envia MP3 para revisão, cria musicSubmissions, lista submissões do user/admin, aprova, rejeita, cancela, completa e pede email de titularidade."],
    ["paths.ts", "Mapa central de coleções Firestore e caminhos Storage. Evita erros de escrever strings repetidas em vários ficheiros."],
    ["profileRequests.ts", "Pedidos ligados ao perfil: mudar displayName, pedir para apagar track/album, reportar perfil, listar pedidos rejeitados/pendentes e admin aprovar/rejeitar."],
    ["releaseReminderClient.ts", "Lembretes de pré-lançamento. Cria ou remove documento releaseReminders com ID userId_albumId."],
    ["schema.ts", "Tipos TypeScript da base de dados. Explica formato de users, albums, tracks, posts, submissions, notifications, reports, eventRequests e appConfig."],
    ["settingsClient.ts", "Configurações do user. Atualiza settings, histórico de pesquisa, bloqueios e biblioteca de albums guardados usando arrayUnion/arrayRemove."],
    ["socialClient.ts", "Ações sociais. Likes e follows por Cloud Functions com idToken; comentários, reports, verificationRequests, notificações e overview admin."],
    ["storageClient.ts", "Uploads. Decide caminho do ficheiro, transforma URI em Blob, envia para Storage com progresso, define contentType e devolve downloadUrl."],
    ["userProfile.ts", "Cria perfil inicial se não existir e atualiza dados do perfil. Garante que Auth user tem documento users/{uid}."],
    ["validate.ts", "Validação simples de email com regex. Usado para conferir se o formato do email parece válido."],
]


FIREBASE_CODE_PATTERNS = [
    ["addDoc(collection(db, nome), dados)", "Cria um documento novo com ID automático. Usado em posts, tracks, albums, submissions, reports, eventRequests."],
    ["doc(db, caminho)", "Aponta para um documento específico. Exemplo: doc(db, firestorePaths.user(userId))."],
    ["getDoc(ref)", "Lê um documento específico uma vez."],
    ["getDocs(query(...))", "Lê vários documentos de uma coleção com filtros/ordenação."],
    ["query(collection(...), where(...), orderBy(...), limit(...))", "Monta uma consulta Firestore. Exemplo: músicas publicadas ordenadas por data."],
    ["updateDoc(ref, dados)", "Atualiza campos de um documento existente."],
    ["serverTimestamp()", "Grava data/hora do servidor Firebase, não do telemóvel."],
    ["httpsCallable(functions, 'nome')", "Chama uma Cloud Function callable pelo app."],
    ["getIdToken(true)", "Pede token atual do Firebase Auth para provar a identidade numa Cloud Function."],
    ["uploadBytesResumable", "Faz upload com progresso para Firebase Storage."],
    ["getDownloadURL", "Depois do upload, gera URL pública/legível conforme regras para mostrar media na app."],
    ["writeBatch / transaction", "Agrupa alterações para manter contadores e documentos consistentes. Nas Functions é muito usado para likes/follows/deleções."],
]


FLOW_TABLE = [
    ["Abrir app", "app/_layout.tsx carrega fonte/splash -> app/index.tsx vê sessão -> sem user mostra login/criar conta -> com user vai para /main/home."],
    ["Login", "login.tsx recolhe email/password -> firebase/auth.ts login() -> Firebase Auth valida -> ensureUserProfile garante users/{uid} -> app vai para home."],
    ["Criar conta", "new-account.tsx pede email -> sendSignupOtpEmail envia código -> verify-mail.tsx confirma -> register cria Auth user -> ensureUserProfile cria perfil."],
    ["Recuperar password", "forgot-password pede email -> Cloud Function envia OTP -> verifyOtp cria resetToken -> new-password chama resetPasswordWithOtp -> Auth troca senha."],
    ["Enviar música", "track.tsx escolhe MP3 -> submitMusicForReview cria musicSubmissions -> upload MP3 para Storage -> status fingerprint_queued/manual_review -> admin decide."],
    ["Aprovar música", "admin/music-reviews usa approveMusicSubmission/allowMusicSubmissionBatch -> trigger notifyMusicSubmissionDecision cria notificação -> user finaliza track."],
    ["Finalizar lançamento", "finish-track carrega submission aprovada -> upload capa/vídeo -> createTrack/createAlbum -> se scheduled, publishScheduledReleases publica quando chegar hora."],
    ["Criar post", "createPost escolhe media -> upload postMedia/overlayMedia -> createPost grava posts -> home/search/profile mostram se status published."],
    ["Dar like", "Tela chama createLike/toggleTrackLike -> socialClient pega idToken -> Cloud Function toggleContentLike faz transação -> cria/remove like e ajusta likesCount."],
    ["Seguir perfil", "profile chama createFollow/removeFollow -> Cloud Function setUserFollow cria/remove follows e subcoleções followers/following -> atualiza contadores."],
    ["Pedir evento", "events/request cria eventRequests -> upload media em events/{id}/banner -> admin aprova -> eventBanners aparece na home até expirar."],
    ["Reportar", "Tela chama createReport -> grava reports status open -> admin responde -> notifyReportDecision cria notificação ao reporter."],
    ["Apagar conta", "settings/menu envia código -> Cloud Function valida OTP -> deleteAccountEverywhere apaga Auth, dados, media e relações."],
]


SIMPLE_BIG_FILE_EXPLANATIONS = [
    [
        "home.tsx",
        "A home é a página inicial da área principal. Pensa nela como a montra da Sonnor: mostra músicas, posts, lançamentos, eventos, pré-lançamentos e recomendações.",
        "1. Importa ferramentas: React, Expo Router, player, Firebase e componentes visuais.\n2. Define tipos como FeaturedPost, MusicItem e HomeBanner para organizar os dados.\n3. Tem funções pequenas para formatar likes, datas, countdown e media.\n4. Usa getHomeContent para buscar dados do Firebase.\n5. Guarda dados com useState: homeData, posts ativos, likes, banners, lembretes e tempo atual.\n6. Usa useEffect para carregar dados quando a tela abre e atualizar contadores.\n7. Usa usePlayer para tocar músicas e filas.\n8. Usa router.push para abrir search, release, track ou post.\n9. Usa createLike e createReport para ações sociais.",
        "Quando a home abre, ela pergunta ao Firebase: 'quais músicas, posts, albums, eventos e perfis devo mostrar?'. Depois transforma esses dados em blocos visuais: carrosséis, cards, grelhas, banners e botões. Se o utilizador toca numa música, chama o player; se toca num post, abre fullscreen; se toca num lançamento, navega para release.",
        "Frase para apresentação: O ficheiro home.tsx é o feed principal da Sonnor. Ele carrega conteúdo do Firebase, organiza em secções visuais e liga ações como tocar música, abrir lançamentos, gostar e reportar."
    ],
    [
        "profile.tsx",
        "O profile é a página de perfil do utilizador/artista. Ele mostra identidade, conteúdo publicado, músicas, albums, posts, loja/merch e ações como seguir, editar ou reportar.",
        "1. Define tipos de dados: TrackItem, AlbumItem, PostItem, MerchProduct e ProfileExtras.\n2. Tem funções de ajuda para converter textos/números, formatar datas, likes e countdowns.\n3. Cria componentes internos como MediaBox, PostGridTile e FullscreenPost.\n4. Usa getProfileContent para buscar dados do perfil.\n5. Guarda muitos estados com useState: profileUser, tracks, albums, posts, followersCount, isFollowing, menus de edição, merch e posts fullscreen.\n6. Usa createFollow/removeFollow para seguir ou deixar de seguir.\n7. Usa uploadUriToStorage para trocar imagens como avatar/banner/merch.\n8. Usa updateUserProfile para gravar mudanças.\n9. Usa router.push para abrir editor, organizar perfil, criar post e settings.",
        "Quando o perfil abre, ele descobre qual utilizador deve mostrar. Se for o próprio utilizador, mostra botões de edição. Se for outro, mostra seguir/reportar. Depois busca no Firebase o documento users e conteúdos ligados: tracks, albums e posts. A tela monta tudo em blocos: topo visual, bio, estatísticas, posts, músicas e loja.",
        "Frase para apresentação: O ficheiro profile.tsx é a tela que transforma os dados do utilizador em uma página pública de artista/ouvinte, com conteúdo, identidade visual e ações de perfil."
    ],
    [
        "search.tsx",
        "O search é a área de pesquisa e descoberta. Ele deixa procurar músicas, artistas, albums e posts, além de explorar categorias e posts aleatórios.",
        "1. Define tipos como ResultItem, CategoryItem, PostItem, RecentItem e DetailState.\n2. Tem funções para normalizar pesquisa, embaralhar posts, formatar likes e tocar clips de post.\n3. Usa getHomeContent e getSearchableUsers para montar resultados.\n4. Guarda query, categoria ativa, histórico, posts, resultados e post fullscreen com useState.\n5. Usa useMemo para filtrar resultados sem recalcular tudo sempre.\n6. Usa router.push para abrir perfil, release, track ou library.\n7. Usa usePlayer para tocar filas de músicas.\n8. Usa createLike e createReport para ações sociais em posts.\n9. Mostra grelhas e fullscreen usando componentes internos.",
        "Quando o utilizador escreve na pesquisa, o texto é normalizado e comparado com títulos, artistas, usernames, músicas relacionadas e categorias. Se não estiver pesquisando, a tela mostra explorar/categorias. Quando clica num resultado, a app decide para onde navegar: perfil, lançamento, música ou biblioteca.",
        "Frase para apresentação: O ficheiro search.tsx é o motor de descoberta da Sonnor. Ele pega dados do Firebase, filtra por texto/categorias e abre os conteúdos certos."
    ],
    [
        "createPost.tsx",
        "O createPost é o editor de publicação. Ele é como uma mini ferramenta de criação: escolhe imagem/vídeo, escreve legenda, adiciona overlays, associa música e publica.",
        "1. Define limites: MAX_CHARS, MAX_VIDEO_DURATION_SECONDS, POST_MUSIC_CLIP_SECONDS, tamanho do palco e zona de lixo.\n2. Define tipos como OverlayItem e ReadySong.\n3. Tem funções para descobrir extensão do ficheiro, duração, tamanho de overlay e limitar escala.\n4. Guarda estados do editor: media base, tipo, escala, overlays, legenda, filtro, música selecionada, clip start/end, preview e publishing.\n5. Usa getHomeContent para listar músicas prontas que podem ser associadas ao post.\n6. Usa uploadUriToStorage para enviar media principal e overlays para Storage.\n7. Usa createPost para criar documento posts no Firestore.\n8. Usa updatePostMedia para completar o post depois dos uploads.\n9. Usa router.back para voltar quando termina.",
        "O fluxo é: o utilizador escolhe media, edita visual, escreve legenda, escolhe música se quiser, define o trecho de áudio e toca em publicar. O código cria primeiro o documento do post, faz upload da media para Storage, guarda os URLs no Firestore e volta para a tela anterior.",
        "Frase para apresentação: O ficheiro createPost.tsx é o editor de posts da Sonnor. Ele junta media, texto, overlays e música, envia ficheiros para Storage e cria o post no Firestore."
    ],
]


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.color.rgb = RGBColor(46, 116, 181)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor(31, 77, 120)
    styles["Heading 3"].font.name = "Arial"
    styles["Heading 3"].font.size = Pt(11)
    styles["Heading 3"].font.color.rgb = RGBColor(31, 77, 120)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Manual completo da app Sonnor")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(15, 15, 15)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Guia para explicar o projeto, o código, o Firebase, a base de dados e o visual na apresentação de 3 de julho de 2026.")
    doc.add_paragraph()

    add_heading(doc, "1. Tree do projeto inteiro", 1)
    doc.add_paragraph("Esta árvore mostra o projeto atual sem pastas gigantes geradas automaticamente, como node_modules, .git, .expo, tmp e caches. Essas pastas existem no computador, mas não são código escrito da app.")
    add_code_block(doc, "Sonnor/\n" + "\n".join(list_tree(ROOT)), max_chars=18000)

    add_heading(doc, "2. Resumo para dizer no início da apresentação", 1)
    doc.add_paragraph(
        "A Sonnor é uma aplicação mobile de música e rede social construída com Expo, React Native, TypeScript e Firebase. "
        "Ela permite criar conta, montar perfil de artista/ouvinte, publicar posts com imagens ou vídeos, enviar músicas para revisão, "
        "finalizar lançamentos, ouvir faixas no player, seguir utilizadores, gostar de conteúdos, guardar albums, pedir eventos, receber notificações e usar um painel administrativo."
    )
    add_bullets(doc, [
        "Frontend: Expo Router organiza as telas por ficheiros dentro de app/.",
        "Estado global: PlayerContext controla áudio, fila, histórico e recomendações.",
        "Backend: Firebase Auth, Firestore, Storage e Cloud Functions.",
        "Base de dados: coleções tipadas em firebase/schema.ts e protegidas por firestore.rules.",
        "Uploads: Storage guarda avatares, banners, capas, áudios MP3, vídeos curtos, posts e eventos.",
        "Admin: telas em app/admin e funções cloud para revisão musical, eventos, verificação, reports e eliminação.",
    ])

    add_heading(doc, "3. Tecnologias e tudo que foi instalado", 1)
    doc.add_paragraph("O package.json mostra as dependências instaladas para a app mobile. O package-lock.json fixa versões exatas para reprodução.")
    add_table(doc, ["Pacote", "Versão", "Para que serve"], package_rows(), [2.2, 1.2, 3.0])

    functions_pkg = json.loads(read_text("functions/package.json"))
    rows = []
    for name, version in functions_pkg["dependencies"].items():
        purpose = {
            "firebase-admin": "SDK administrativo do Firebase usado no backend. Permite às Cloud Functions ler/escrever Firestore, gerir Auth, apagar utilizadores, mexer em Storage e criar timestamps do servidor.",
            "firebase-functions": "Framework das Cloud Functions. É o que permite criar funções callable chamadas pela app, endpoints HTTP, tarefas agendadas e triggers quando documentos Firestore mudam.",
            "nodemailer": "Biblioteca de envio de email. Na Sonnor envia códigos OTP de signup/password/delete e emails de prova de titularidade musical através de Gmail.",
        }[name]
        rows.append([name, version, purpose])
    for name, version in functions_pkg["devDependencies"].items():
        purpose = {
            "@types/node": "Tipos TypeScript para Node.js dentro das Cloud Functions. Ajuda o editor/compilador a entender crypto, path, process e APIs do servidor.",
            "firebase-functions-test": "Ferramenta de testes para Cloud Functions. Serve para simular funções/triggers em desenvolvimento, mesmo que não exista uma suíte grande no projeto.",
            "typescript": "Compilador TypeScript das Functions. Transforma functions/src/index.ts em JavaScript dentro de functions/lib/index.js antes do deploy.",
        }[name]
        rows.append([name, version, purpose])
    add_heading(doc, "Dependências das Cloud Functions", 2)
    add_table(doc, ["Pacote", "Versão", "Para que serve"], rows, [2.2, 1.2, 3.0])

    add_heading(doc, "4. Configuração do projeto", 1)
    add_table(doc, ["Ficheiro", "Explicação"], [
        ["app.json", "Configura nome Sonnor, slug, splash preto, ícones, permissões de galeria, plugins Expo, typed routes e React Compiler."],
        ["package.json", "Scripts para start/android/ios/web/lint e dependências da app."],
        ["tsconfig.json", "Configuração TypeScript usada pelo Expo."],
        ["firebase.json", "Liga Firestore rules, indexes, Storage rules e Functions com predeploy npm run build."],
        [".firebaserc", "Projeto Firebase default: sonnor-d8a30."],
        ["firestore.rules", "Regras de segurança de leitura/escrita da base de dados."],
        ["storage.rules", "Regras de segurança dos ficheiros no Storage."],
        ["firestore.indexes.json", "Índices necessários para queries com where/orderBy."],
    ], [2.1, 4.3])

    add_heading(doc, "5. Como a app abre", 1)
    add_numbered(doc, [
        "Expo entra por expo-router/entry, definido no package.json.",
        "O router carrega app/_layout.tsx.",
        "app/_layout.tsx carrega a fonte Bristol, mantém o splash nativo, mostra um ecrã preto com a letra S e prepara Firebase.",
        "Enquanto isso, escuta o estado do Firebase Auth. Se houver utilizador, faz um pré-carregamento de home com getHomeContent.",
        "Quando fonte, animação e Firebase estão prontos, esconde a abertura e renderiza as rotas.",
        "A app fica envolvida por PlayerProvider e SuccessFeedbackProvider, então qualquer tela pode usar o player e feedback visual.",
    ])

    add_heading(doc, "6. Telas e rotas", 1)
    add_table(doc, ["Ficheiro", "Tela", "O que faz"], ROUTE_EXPLANATIONS, [2.05, 1.35, 3.05])

    add_heading(doc, "7. Visual e experiência", 1)
    doc.add_paragraph(
        "O visual da Sonnor é escuro, musical e social. A base é preto/cinza, com cards de media, imagens grandes, vídeos, blur, cantos arredondados, "
        "botões com feedback de pressão e uma identidade de abertura com a fonte Bristol. A app usa conteúdos visuais reais: capas, banners, posts, vídeos e GIFs/assets."
    )
    add_bullets(doc, [
        "Splash/branding: app/_layout.tsx mostra uma abertura preta com a letra S em Bristol.",
        "Estrutura de ecrã: components/AppScreen.tsx centraliza safe areas, teclado e fundo escuro.",
        "Menu: app/main/components/Dynamicmenu.tsx cria navegação inferior e player mini, com muitas animações e estados.",
        "Media fullscreen: FullscreenMedia.tsx e fullmidia.tsx tratam imagens/vídeos em ecrã inteiro.",
        "Criação de post: createPost.tsx tem editor visual com seleção de media, overlays, escala, ligação a música e clip range.",
        "Player: LinearSeekBar.tsx, ClipRangeSelector.tsx e SharedMediaProgress.tsx cuidam de progresso e cortes de áudio.",
        "Feedback: SuccessFeedback.tsx cria feedback global para ações bem-sucedidas.",
    ])

    add_heading(doc, "8. Player de música", 1)
    doc.add_paragraph(
        "O player vive em context/PlayerContext.tsx. Ele é uma camada global para tocar música a partir da home, search, perfil, release ou biblioteca."
    )
    add_bullets(doc, [
        "Track type guarda id, uri, title, artist, cover, shortVideo, lyrics, albumId, folderTitle, genre e origem.",
        "playTrack toca uma faixa isolada e limpa fila.",
        "playQueue toca uma lista de músicas a partir de um índice.",
        "playNext e playPrevious controlam fila e histórico.",
        "autoRecommendations permite continuar a tocar recomendações quando a fila acaba.",
        "prepareSound pré-carrega a próxima faixa para diminuir espera.",
        "canPlayAlbumTrack impede tocar faixas de pré-lançamento antes da hora.",
        "createRecentPlay grava histórico em recentPlays quando há utilizador autenticado.",
    ])

    add_heading(doc, "9. Firebase: projeto, Auth, Firestore, Storage e Functions", 1)
    add_table(doc, ["Parte", "Detalhes"], [
        ["Projeto Firebase", "sonnor-d8a30, definido em .firebaserc e firebase/config.ts."],
        ["Auth", "Email/senha com persistência local via AsyncStorage. Fluxos customizados de OTP por Cloud Functions."],
        ["Firestore", "Base de dados NoSQL com coleções para users, tracks, albums, posts, submissions, likes, follows, reports, eventos e configs."],
        ["Storage", "Ficheiros de media: avatares, banners, capas, áudios, vídeos, posts, eventos e uploads temporários."],
        ["Functions", "Backend Node 20 com callable functions, HTTP fallbacks, triggers Firestore e schedules."],
        ["Secrets", "GMAIL_USER e GMAIL_APP_PASSWORD usados por nodemailer nas Functions. Isto não fica no app.json; é segredo do ambiente Firebase."],
    ], [1.8, 4.7])

    add_heading(doc, "10. Base de dados Firestore", 1)
    doc.add_paragraph("A base de dados está documentada em firebase/schema.ts, os nomes/caminhos em firebase/paths.ts e as permissões em firestore.rules.")
    add_table(doc, ["Coleção", "Responsabilidade", "Campos principais"], DATA_MODEL, [1.55, 1.8, 3.1])

    add_heading(doc, "11. Regras de segurança da base de dados", 1)
    add_bullets(doc, [
        "users pode ser lido publicamente, mas só o próprio utilizador edita campos permitidos; admin pode mais.",
        "albums/tracks/posts só são públicos quando published ou scheduled conforme a regra; dono e admin também podem ler.",
        "tracks só podem ser criadas pelo utilizador quando existe musicSubmission aprovada com o mesmo áudio, ou por admin.",
        "likes e follows não são escritos diretamente pelo cliente; são controlados por Cloud Functions para evitar contadores falsos.",
        "reports são criados pelo utilizador, mas lidos e geridos por admin.",
        "eventRequests e verificationRequests são criados pelo dono e decididos por admin.",
        "appConfig/public pode ser lido por todos; appConfig/admins só por admin.",
    ])

    add_heading(doc, "12. Storage e organização de ficheiros", 1)
    add_table(doc, ["Caminho", "Uso"], [
        ["users/{userId}/avatar, banner, background", "Imagens do perfil."],
        ["albums/{albumId}/cover.jpg e background.jpg", "Capa e fundo de lançamentos."],
        ["tracks/{trackId}/audio.mp3, preview.mp3, short-video.mp4, cover.jpg", "Áudio, prévia, vídeo curto e capa de faixas."],
        ["musicSubmissions/{submissionId}/source.mp3", "MP3 enviado para revisão antes de publicar."],
        ["posts/{postId}/media e overlays", "Imagem/vídeo do post e imagens sobrepostas."],
        ["events/{eventId}/banner", "Media de eventos enviados e aprovados."],
        ["messages/{threadId}/{messageId}", "Ficheiros de mensagens, estrutura preparada."],
        ["tempUploads/{userId}", "Uploads temporários."],
    ], [2.6, 3.9])
    doc.add_paragraph("storage.rules valida dono, admin, participante de thread, tipo de ficheiro e tamanho máximo em alguns caminhos: imagens até 12 MB, vídeos até 80 MB e MP3 de submissão até 60 MB.")

    add_heading(doc, "13. Cloud Functions", 1)
    add_table(doc, ["Função", "Responsabilidade"], CLOUD_FUNCTIONS, [2.7, 3.8])

    add_heading(doc, "14. Clientes Firebase no app", 1)
    add_table(doc, ["Ficheiro", "Responsabilidade"], CLIENT_EXPLANATIONS, [2.3, 4.2])

    add_heading(doc, "15. Fluxos principais para explicar", 1)
    add_heading(doc, "Criar conta", 2)
    add_numbered(doc, [
        "O utilizador escreve email em new-account.tsx.",
        "firebase/auth.ts chama sendSignupOtpEmail.",
        "A Cloud Function gera código de 4 dígitos, guarda hash em signupOtps e envia email.",
        "verify-mail.tsx confirma o código com verifySignupOtp.",
        "Depois é criado o utilizador no Firebase Auth.",
        "ensureUserProfile cria o documento users/{uid} com contadores a zero.",
    ])
    add_heading(doc, "Enviar e publicar música", 2)
    add_numbered(doc, [
        "track.tsx escolhe MP3 e envia para Storage em musicSubmissions/{id}/source.mp3.",
        "Cria musicSubmissions com estado uploaded/fingerprint/manual review e declaração de direitos.",
        "Admin usa music-reviews.tsx para aprovar, rejeitar ou pedir prova.",
        "Quando aprovado, notifyMusicSubmissionDecision envia notificação.",
        "finish-track/[submissionId].tsx permite completar título, capa, letra, género, vídeo curto e lançamento.",
        "createTrack/createAlbum gravam tracks/albums. Se houver data futura, fica scheduled; quando chega a hora, publishScheduledReleases publica.",
    ])
    add_heading(doc, "Criar post", 2)
    add_numbered(doc, [
        "createPost.tsx abre editor de imagem/vídeo.",
        "Pode associar uma música já pronta e escolher o corte do clip.",
        "Uploads vão para posts/{postId}/media e overlays.",
        "createPost cria documento posts com caption, mediaUrl, linkedTrackId e contadores a zero.",
        "Home/search/profile leem posts published e mostram em feed ou grelha.",
    ])
    add_heading(doc, "Gostar e seguir", 2)
    add_numbered(doc, [
        "A interface chama socialClient.ts.",
        "socialClient chama Cloud Functions com idToken.",
        "A função faz transação no Firestore: cria/remove like ou follow e atualiza contadores.",
        "Isto evita que o cliente manipule likesCount/followersCount diretamente.",
    ])
    add_heading(doc, "Eventos", 2)
    add_numbered(doc, [
        "events/request.tsx permite enviar título, detalhes, link e imagem/vídeo.",
        "O pedido fica em eventRequests com status pending.",
        "Admin em event-banners.tsx aprova/rejeita.",
        "Quando aprovado, cria eventBanners publicado por uma semana.",
        "Home mostra apenas banners publicados e ainda não expirados; banners followers aparecem a seguidores.",
    ])

    add_heading(doc, "16. Painel admin", 1)
    add_bullets(doc, [
        "app/admin/index.tsx: dashboard e atalhos principais.",
        "music-reviews.tsx: revisão de submissões MP3, direitos, aprovação/rejeição e emails de prova.",
        "manage-users.tsx: gestão de utilizadores, verificação e ações sensíveis.",
        "user-requests.tsx: pedidos feitos pelos utilizadores.",
        "event-banners.tsx: aprovação e publicação de banners de eventos.",
        "manage-releases.tsx, manage-posts.tsx e manage-reports.tsx: páginas de gestão/atalhos administrativos.",
        "verification-requests.tsx: entrada para pedidos de verificação.",
    ])

    add_heading(doc, "17. Interface studio-interface", 1)
    doc.add_paragraph(
        "A pasta studio-interface é uma interface web separada da app mobile. Tem index.html, style.css, app.js e server.js. "
        "Ela serve como ferramenta local/auxiliar para visualizar ou gerir informação sem entrar pela app Expo. O server.js cria o servidor local, enquanto app.js contém a lógica do browser."
    )

    add_heading(doc, "18. Componentes, hooks, utils e assets", 1)
    add_table(doc, ["Área", "Ficheiros", "Explicação"], [
        ["components", "AppScreen, ClipRangeSelector, LinearSeekBar, SuccessFeedback, pressFeedback", "Blocos reutilizáveis para layout, sliders, range de clips, feedback e animação de toque."],
        ["hooks", "useCurrentUser, useAsyncData", "Pequenas abstrações para utilizador atual e carregamento assíncrono."],
        ["utils", "mediaPicker, uploadAsset, responsive, avatarFallback", "Escolha de media, uploads, medidas responsivas e cor de avatar quando não há imagem."],
        ["icons", "BackIcon, EyeOpen/EyeClosed, Key, Lock, Mail", "Ícones usados nas telas de auth e navegação."],
        ["assets", "GIFs, MP4, imagens de ícone/splash, Bristol.ttf", "Materiais visuais e identidade da aplicação."],
        ["constants", "musicGenres, musicLibrary", "Listas estáticas de géneros e conteúdo/base auxiliar."],
    ], [1.2, 2.35, 3.0])

    add_heading(doc, "19. Aula: como as screens são chamadas", 1)
    doc.add_paragraph(
        "A Sonnor usa Expo Router. Isso significa que a pasta app/ funciona como mapa de navegação. "
        "O nome do ficheiro vira o caminho da tela. Para abrir uma tela, o código usa router.push, router.replace ou router.back."
    )
    add_table(doc, ["Ficheiro", "Como funciona como rota"], SCREEN_ROUTING_LESSONS, [2.5, 4.0])
    add_heading(doc, "Exemplo de navegação", 2)
    add_code_block(doc, """const router = useRouter();

router.push("/main/profile");
router.replace("/auth/login");
router.back();""")
    add_bullets(doc, [
        "router.push('/main/profile') abre a tela de perfil e mantém histórico para voltar.",
        "router.replace('/auth/login') troca a tela atual; é usado quando não faz sentido voltar para a anterior.",
        "router.back() volta para a tela anterior.",
        "Rotas com [id], [slug] ou [submissionId] recebem valores pela URL.",
    ])

    add_heading(doc, "20. Aula: como se cria uma screen nova", 1)
    add_numbered(doc, SCREEN_CREATION_STEPS)
    add_heading(doc, "Modelo simples de uma screen", 2)
    add_code_block(doc, """import { useRouter } from "expo-router";
import { Pressable, StyleSheet, Text, View } from "react-native";

export default function MinhaScreen() {
  const router = useRouter();

  return (
    <View style={styles.container}>
      <Text style={styles.title}>Titulo</Text>
      <Pressable style={styles.button} onPress={() => router.push("/main/home")}>
        <Text style={styles.buttonText}>Ir para Home</Text>
      </Pressable>
    </View>
  );
}

const styles = StyleSheet.create({
  container: { flex: 1, backgroundColor: "#000", padding: 24 },
  title: { color: "#fff", fontSize: 28, fontWeight: "800" },
  button: { backgroundColor: "#fff", borderRadius: 24, padding: 14 },
  buttonText: { color: "#000", textAlign: "center", fontWeight: "700" },
});""")

    add_heading(doc, "21. Aula: blocos visuais usados na app", 1)
    doc.add_paragraph(
        "No React Native, o visual é montado com blocos. Uma tela é uma árvore de componentes: View dentro de View, Text dentro de View, Pressable envolvendo botões, Image mostrando media, e assim por diante."
    )
    add_table(doc, ["Bloco", "Para que serve"], VISUAL_BUILDING_BLOCKS, [1.7, 4.8])
    add_heading(doc, "Exemplo de card/bloco", 2)
    add_code_block(doc, """<View style={styles.card}>
  <Text style={styles.cardTitle}>Nova musica</Text>
  <Text style={styles.cardDescription}>Descricao do bloco.</Text>
</View>

const styles = StyleSheet.create({
  card: {
    borderRadius: 22,
    padding: 18,
    backgroundColor: "rgba(255,255,255,0.055)",
    borderWidth: 1,
    borderColor: "rgba(255,255,255,0.075)",
  },
  cardTitle: { color: "#fff", fontSize: 18, fontWeight: "800" },
  cardDescription: { color: "#c7c7c7", fontSize: 14, marginTop: 8 },
});""")
    add_bullets(doc, [
        "borderRadius arredonda o bloco.",
        "padding cria espaço interno.",
        "backgroundColor define a cor do fundo.",
        "borderWidth e borderColor criam contorno sutil.",
        "fontSize, color e fontWeight controlam texto.",
    ])

    add_heading(doc, "22. Aula: como o menu foi criado", 1)
    add_table(doc, ["Parte", "Explicação"], MENU_LESSONS, [2.0, 4.5])
    add_heading(doc, "Código principal que decide se o menu aparece", 2)
    add_code_block(doc, """const pathname = usePathname();

const showDynamic =
  pathname === "/main/home" ||
  pathname === "/main/search" ||
  pathname === "/main/profile" ||
  pathname === "/main/library" ||
  pathname.startsWith("/main/release/");

return (
  <View style={{ flex: 1, backgroundColor: "#000" }}>
    <Stack screenOptions={{ headerShown: false, animation: "fade" }} />
    {showDynamic && <Dynamicmenu />}
  </View>
);""")
    add_bullets(doc, [
        "usePathname lê a rota atual.",
        "showDynamic fica true apenas em telas onde queremos o menu.",
        "Stack mostra a screen atual.",
        "{showDynamic && <Dynamicmenu />} significa: se showDynamic for true, mostra o menu.",
    ])

    add_heading(doc, "23. Aula: padrões visuais da Sonnor", 1)
    add_table(doc, ["Padrão visual", "Como foi feito"], VISUAL_PATTERNS, [1.8, 4.7])
    doc.add_paragraph(
        "Quando te perguntarem sobre o visual, podes dizer: o design foi montado com React Native usando View, Text, Image, Pressable, Modal, BlurView, StyleSheet e componentes próprios. "
        "A identidade usa fundo preto, media grande, cards translúcidos, bordas arredondadas, ícones e animações."
    )

    add_heading(doc, "24. Aula: ficheiros mais importantes do código", 1)
    add_table(doc, ["Ficheiro", "O que deves saber explicar"], CODE_FILE_LESSONS, [2.4, 4.1])

    add_heading(doc, "25. Aula: telas explicadas com mais profundidade", 1)
    doc.add_paragraph(
        "Esta parte é para quando te perguntarem uma tela específica. Em vez de responder só 'essa tela mostra posts', podes explicar o que ela carrega, que ficheiro usa e que ações faz."
    )
    add_table(doc, ["Tela/ficheiro", "Explicação mais profunda"], DEEP_SCREEN_LESSONS, [2.5, 4.0])

    add_heading(doc, "26. Explicação fácil dos ficheiros grandes", 1)
    doc.add_paragraph(
        "Estes quatro ficheiros são grandes porque juntam muito visual e muita lógica. Para a apresentação, não precisas decorar linha por linha. "
        "Precisas saber explicar o papel de cada ficheiro, como recebe dados, como monta o visual e que ações executa."
    )
    for file_name, simple_role, parts, flow, speech in SIMPLE_BIG_FILE_EXPLANATIONS:
        add_heading(doc, file_name, 2)
        doc.add_paragraph(simple_role)
        add_table(doc, ["Parte", "Explicação"], [
            ["Partes principais", parts],
            ["Como funciona", flow],
            ["Como dizer na apresentação", speech],
        ], [1.6, 4.9])

    add_heading(doc, "27. Aula: Firebase no código do app", 1)
    doc.add_paragraph(
        "A pasta firebase/ é a ponte entre as telas e o backend. As telas não falam diretamente com tudo de forma solta; elas chamam funções organizadas por responsabilidade."
    )
    add_table(doc, ["Ficheiro Firebase", "Como funciona no código"], FIREBASE_CLIENT_DEEP, [2.35, 4.15])

    add_heading(doc, "Todos os ficheiros da pasta firebase, um por um", 2)
    doc.add_paragraph(
        "Esta tabela fecha a pasta firebase inteira. Se perguntarem qualquer ficheiro dessa pasta, podes usar esta explicação curta."
    )
    add_table(doc, ["Ficheiro", "Para que serve"], ALL_FIREBASE_FILES, [2.0, 4.5])

    add_heading(doc, "28. Aula: padrões de código Firebase", 1)
    add_table(doc, ["Padrão", "Significado na Sonnor"], FIREBASE_CODE_PATTERNS, [2.45, 4.05])
    add_heading(doc, "Exemplo Firebase explicado", 2)
    add_code_block(doc, """const snapshot = await getDocs(
  query(
    collection(db, firestoreCollections.posts),
    where("status", "==", "published"),
    orderBy("createdAt", "desc"),
    limit(18),
  ),
);

const posts = snapshot.docs.map((docSnap) => ({
  id: docSnap.id,
  ...docSnap.data(),
}));""")
    add_bullets(doc, [
        "collection(db, firestoreCollections.posts) escolhe a coleção posts.",
        "where('status', '==', 'published') filtra só posts publicados.",
        "orderBy('createdAt', 'desc') coloca os mais recentes primeiro.",
        "limit(18) limita a quantidade para performance.",
        "getDocs executa a query.",
        "snapshot.docs.map transforma documentos Firestore em objetos normais com id e dados.",
    ])

    add_heading(doc, "29. Aula: fluxos completos da app", 1)
    add_table(doc, ["Fluxo", "Passo a passo"], FLOW_TABLE, [1.7, 4.8])

    add_heading(doc, "30. Aula: Cloud Functions por dentro", 1)
    doc.add_paragraph(
        "As Cloud Functions existem porque algumas ações não devem depender apenas do telemóvel. O backend confirma identidade, protege regras, atualiza contadores e executa tarefas automáticas."
    )
    add_bullets(doc, [
        "OTP: gera código, guarda apenas hash no Firestore, define expiração e limita tentativas.",
        "Emails: usa nodemailer com secrets GMAIL_USER e GMAIL_APP_PASSWORD para enviar códigos e pedidos de prova.",
        "Likes/follows: usa transações para impedir contadores errados quando duas ações acontecem ao mesmo tempo.",
        "Publicação agendada: publishScheduledReleases corre a cada minuto e publica albums/faixas quando releaseDate chega.",
        "Eventos: expireEventBanners corre diariamente e marca banners expirados.",
        "Notificações: triggers onDocumentWritten criam notificações quando admin aprova/rejeita pedidos.",
        "Verificação: recalcula verified automaticamente por seguidores, músicas com muitos likes e reports aprovados.",
        "Deleção: deleteAccountEverywhere limpa dados ligados para não deixar conteúdo órfão.",
    ])

    add_heading(doc, "31. Índices Firestore", 1)
    doc.add_paragraph(
        "firestore.indexes.json define índices compostos para queries como status + createdAt, userId + status, releaseDate, notifications por userId e musicSubmissions por status. "
        "Sem esses índices, algumas listas da home, perfil, biblioteca e admin falhariam ou ficariam lentas."
    )

    add_heading(doc, "32. Git e GitHub", 1)
    doc.add_paragraph(
        "Git é o sistema que guarda o histórico do código. GitHub é o site onde esse histórico pode ficar online. "
        "Na Sonnor, o repositório local está conectado ao GitHub pelo remote origin."
    )
    add_table(doc, ["Item", "Valor no projeto Sonnor"], [
        ["Remote origin", "https://github.com/zepriquito7k/Sonnor"],
        ["Branch atual", "main"],
        ["Últimos commits vistos", "c6995ef Firebase Semi-Pronto; 92da6e6 2.0; e3cfa0a New Structure; 40a36f7 Estrutura MAIN e AUTH feita; 0921522 Estrutura auth."],
        ["O que é commit", "Um ponto salvo no histórico do projeto, com uma mensagem dizendo o que mudou."],
        ["O que é branch", "Uma linha de trabalho separada. A main costuma ser a versão principal."],
        ["O que é remote", "O endereço do repositório online, normalmente no GitHub."],
    ], [2.0, 4.5])
    doc.add_paragraph("Frase para apresentação: o projeto Sonnor tem controlo de versões com Git e está ligado a um repositório GitHub chamado Sonnor. Isso permite guardar versões, voltar atrás se necessário e enviar o código para a nuvem.")
    add_heading(doc, "Comandos Git/GitHub importantes", 2)
    add_table(doc, ["Comando", "Para que serve"], GIT_COMMANDS, [2.25, 4.25])
    add_heading(doc, "Fluxo normal para guardar alterações", 2)
    add_numbered(doc, [
        "Ver o que mudou: git status.",
        "Ver detalhes do que mudou: git diff.",
        "Preparar ficheiros: git add nome-do-ficheiro ou git add .",
        "Criar commit: git commit -m \"explica o que foi feito\".",
        "Enviar para GitHub: git push.",
        "Se outra pessoa mexeu no GitHub antes: git pull para baixar antes de continuar.",
    ])
    doc.add_paragraph(
        "Cuidado importante: nunca colocar passwords, chaves privadas ou ficheiros secretos no GitHub. "
        "No projeto, as credenciais sensíveis de email das Cloud Functions são secrets do Firebase: GMAIL_USER e GMAIL_APP_PASSWORD."
    )

    add_heading(doc, "33. Como entender o código TypeScript/React Native", 1)
    doc.add_paragraph(
        "A Sonnor usa TypeScript, que é JavaScript com tipos. O código aparece muitas vezes com const, async, await, type, useState e useEffect. "
        "Essas palavras não são aleatórias: elas organizam dados, telas, funções e chamadas ao Firebase."
    )
    add_table(doc, ["Código/palavra", "Explicação simples"], CODE_BASICS, [1.65, 4.85])
    add_heading(doc, "Exemplo explicado", 2)
    add_code_block(doc, """const [loading, setLoading] = useState(false);

async function handleLogin() {
  try {
    setLoading(true);
    await login(email, password);
    router.replace("/main/home");
  } catch (error) {
    Alert.alert("Erro", "Nao foi possivel entrar.");
  } finally {
    setLoading(false);
  }
}""")
    add_bullets(doc, [
        "const [loading, setLoading] cria um estado: loading guarda o valor atual e setLoading muda esse valor.",
        "useState(false) começa com loading falso.",
        "async function handleLogin() define uma função que pode esperar operações como login no Firebase.",
        "try/catch/finally tenta entrar, mostra erro se falhar e desliga o loading no final.",
        "await login(email, password) espera o Firebase responder antes de navegar.",
        "router.replace('/main/home') troca a tela atual pela home.",
    ])

    add_heading(doc, "34. Como o código inteiro se encaixa", 1)
    add_numbered(doc, [
        "O utilizador toca na interface em uma tela dentro de app/.",
        "A tela chama uma função cliente em firebase/, utils/ ou context/.",
        "Se for leitura simples, firebase/contentClient.ts lê Firestore e devolve dados.",
        "Se for upload, firebase/storageClient.ts manda ficheiros para Storage e devolve URL.",
        "Se for ação sensível, firebase/socialClient.ts ou auth.ts chama uma Cloud Function.",
        "A Cloud Function valida quem é o utilizador, mexe no Firestore/Auth/Storage e devolve resposta.",
        "A tela atualiza estado com useState/useEffect e mostra novo conteúdo.",
        "O PlayerContext fica por cima das telas para a música continuar enquanto o utilizador navega.",
    ])
    add_heading(doc, "Padrão que se repete em quase todas as telas", 2)
    add_bullets(doc, [
        "Imports no topo: trazem React, Expo, Firebase, componentes e tipos.",
        "Types: descrevem formato dos dados usados pela tela.",
        "useState/useRef/useEffect: guardam estado, referências e carregamentos.",
        "Funções handle...: respondem a ações do utilizador, como publicar, guardar, tocar ou enviar.",
        "return: desenha a interface com View, Text, Pressable, Image, ScrollView e outros componentes.",
        "StyleSheet.create: concentra cores, tamanhos, espaçamentos e layout visual.",
    ])

    add_heading(doc, "35. Comandos úteis", 1)
    add_table(doc, ["Comando", "Quando usar"], [
        ["npm install", "Instalar dependências da app depois de clonar/abrir o projeto."],
        ["npm run start", "Abrir Expo e escolher Android/iOS/Web."],
        ["npm run android", "Abrir direto em Android."],
        ["npm run ios", "Abrir direto em iOS/simulador."],
        ["npm run web", "Testar no navegador."],
        ["npm run lint", "Verificar erros de estilo/código."],
        ["cd functions && npm run build", "Compilar Cloud Functions TypeScript."],
        ["firebase deploy --only firestore:rules,storage,functions", "Publicar regras e funções no Firebase."],
    ], [2.4, 4.1])

    add_heading(doc, "36. Como defender tecnicamente", 1)
    add_bullets(doc, [
        "Começa dizendo que a app está dividida em três camadas: interface Expo, clientes Firebase e backend Firebase.",
        "Explica que a segurança não depende só da interface: firestore.rules e storage.rules impedem alterações indevidas.",
        "Realça que likes/follows e ações sensíveis usam Cloud Functions para consistência e proteção.",
        "Mostra que os dados têm tipos em schema.ts, caminhos centralizados em paths.ts e queries com índices.",
        "Explica que o player é global, não preso a uma tela, por isso continua a funcionar ao navegar.",
        "No visual, fala da identidade escura, da fonte Bristol, de media fullscreen, blur, posts visuais e experiência musical.",
        "Para Firebase, lembra: Auth identifica, Firestore guarda documentos, Storage guarda ficheiros, Functions executa lógica segura.",
        "Para GitHub, explica que Git guarda versões locais e GitHub guarda o repositório online ligado por origin.",
        "Para código, explica que const guarda valores, funções executam ações, hooks guardam estado e Firebase guarda/consulta dados.",
    ])

    add_heading(doc, "37. Mapa de ficheiros por tamanho", 1)
    files = []
    for path in ROOT.rglob("*"):
        if path.is_file() and path.suffix in {".ts", ".tsx", ".js", ".json"}:
            if any(part in SKIP_DIRS for part in path.relative_to(ROOT).parts):
                continue
            files.append([rel(path), str(line_count(path))])
    files = sorted(files, key=lambda row: row[0].lower())
    add_table(doc, ["Ficheiro", "Linhas"], files, [5.4, 0.8])

    add_heading(doc, "38. Mini glossário do código", 1)
    add_table(doc, ["Termo", "Significado"], [
        ["status: draft/scheduled/published", "Controla se conteúdo está em rascunho, agendado ou público."],
        ["slug", "Texto amigável para URL/rota, usado em releases."],
        ["sourceSubmissionId", "Liga uma track final à submissão musical aprovada."],
        ["idToken", "Token do Firebase Auth enviado para Functions verificarem identidade."],
        ["callable function", "Função chamada diretamente pelo app via Firebase Functions."],
        ["onDocumentWritten", "Trigger que executa quando um documento Firestore muda."],
        ["serverTimestamp", "Data gerada pelo servidor Firebase, mais confiável que data do telemóvel."],
        ["cache bust", "Parametro extra no URL para forçar atualização de imagem após upload."],
        ["Git", "Sistema que guarda o histórico do código no computador."],
        ["GitHub", "Site onde o repositório Git fica online."],
        ["commit", "Uma versão salva do projeto."],
        ["push", "Enviar commits do computador para o GitHub."],
        ["pull", "Baixar alterações do GitHub para o computador."],
        ["branch", "Linha separada de desenvolvimento."],
    ], [1.8, 4.7])

    add_heading(doc, "39. Resumo final em 60 segundos", 1)
    doc.add_paragraph(
        "A Sonnor foi construída como uma plataforma musical/social. No telemóvel, o Expo Router organiza telas de auth, home, search, biblioteca, perfil, criação, eventos e admin. "
        "O Firebase Auth cuida da identidade; o Firestore guarda utilizadores, músicas, lançamentos, posts, likes, follows, notificações, reports e eventos; o Storage guarda todos os ficheiros; "
        "e as Cloud Functions protegem ações importantes como OTP, likes, follows, publicação agendada, notificações, revisão musical e eliminação de contas. "
        "O visual foi pensado para música: fundo escuro, media grande, player global, vídeos/imagens fullscreen, capa/fundo dos lançamentos e uma abertura de marca com a fonte Bristol."
    )

    doc.core_properties.title = "Manual completo da app Sonnor"
    doc.core_properties.subject = "Guia tecnico e de apresentacao"
    doc.core_properties.author = "Sonnor"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
